"""
main.py — QP Generation Engine Backend (FastAPI)

Changes from v1:
  - Per-subject VectorStore and chunks (keyed by subject_id)
  - Materials reloaded from DB on startup (persistent across sessions)
  - Generation reads bloom_level, difficulty, question_type from each section pattern
  - Subject CRUD endpoints
  - Paper history endpoints
  - Material listing per subject
"""

from fastapi import FastAPI, UploadFile, File, Form, Depends, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from typing import Dict, Optional
import uuid
import json
import os
import random

# ── RAG ──
from vector_store import VectorStore
from chunker import chunk_text
from data_loader import load_file
from prompt_builder import build_prompt
from llm_engine import generate_question
from query_generator import QueryGenerator
from similarity_checker import is_similar

# ── AUTH ──
from auth import (
    verify_password,
    create_access_token,
    get_current_user,
    require_role,
    JWT_COOKIE_NAME,
)

# ── DB ──
from db import (
    init_db,
    get_user_by_email,
    # subjects
    create_subject,
    get_all_subjects,
    get_subject_by_id,
    delete_subject,
    # materials
    insert_material,
    get_materials_by_subject,
    get_all_materials,
    delete_material,
    # papers
    create_paper,
    get_papers_by_subject,
    get_paper_with_questions,
    # questions
    insert_question,
    update_question_text,
    delete_question_by_id,
    # analytics
    get_analytics_summary,
    get_bloom_distribution,
    get_difficulty_distribution,
)

# ════════════════════════════════════════════════
# PER-SUBJECT IN-MEMORY STORES
# subject_stores[subject_id] = VectorStore instance
# subject_chunks[subject_id] = list of text chunks
# ════════════════════════════════════════════════
subject_stores: Dict[str, VectorStore] = {}
subject_chunks: Dict[str, list] = {}
query_generator = QueryGenerator()

UPLOAD_DIR = "data/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════

def _index_chunks_for_subject(subject_id: str, chunks: list):
    """Add chunks to a subject's in-memory vector store."""
    if subject_id not in subject_chunks:
        subject_chunks[subject_id] = []
    subject_chunks[subject_id].extend(chunks)

    store = VectorStore()
    store.build(subject_chunks[subject_id])
    subject_stores[subject_id] = store


# ════════════════════════════════════════════════
# LIFESPAN — must be defined BEFORE FastAPI()
# ════════════════════════════════════════════════

from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app: FastAPI):
    # ── STARTUP ──
    init_db()
    all_mats = get_all_materials()
    loaded, skipped = 0, 0

    for mat in all_mats:
        fp = mat["file_path"]
        sid = mat["subject_id"]
        if os.path.exists(fp):
            try:
                content = load_file(fp)
                chunks = chunk_text(content)
                if chunks:
                    if sid not in subject_chunks:
                        subject_chunks[sid] = []
                    subject_chunks[sid].extend(chunks)
                    loaded += 1
            except Exception as e:
                print(f"[STARTUP] Failed to load {fp}: {e}")
                skipped += 1
        else:
            print(f"[STARTUP] File not found (skipped): {fp}")
            skipped += 1

    for sid, chunks in subject_chunks.items():
        if chunks:
            store = VectorStore()
            store.build(chunks)
            subject_stores[sid] = store

    print(f"[APP] Startup complete — {loaded} material(s) reloaded, {skipped} skipped.")

    yield  # app runs here
    # ── SHUTDOWN (add cleanup here if needed) ──


# ════════════════════════════════════════════════
# FASTAPI APP  (defined after lifespan)
# ════════════════════════════════════════════════

app = FastAPI(title="QP Generation Engine Backend", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
      "http://localhost:5173",
      "http://127.0.0.1:5173",
      "https://your-vercel-app.vercel.app"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ════════════════════════════════════════════════
# PYDANTIC MODELS
# ════════════════════════════════════════════════

class LoginRequest(BaseModel):
    email: str
    password: str
    role: str

class CreateSubjectRequest(BaseModel):
    name: str
    description: Optional[str] = ""

class GeneratePaperRequest(BaseModel):
    subject_id: str
    title: Optional[str] = "Untitled Paper"
    pattern: Dict               # { sections: [...] }

class EditQuestionRequest(BaseModel):
    question_id: str
    text: str

class RegenerateRequest(BaseModel):
    question_id: str
    subject_id: str
    question_type: Optional[str] = "Short Answer"
    bloom_level: Optional[str] = "Understand"
    difficulty: Optional[str] = "Medium"
    marks: Optional[int] = 2


# ════════════════════════════════════════════════
# AUTH
# ════════════════════════════════════════════════

@app.post("/auth/login")
async def login(req: LoginRequest):
    db_user = get_user_by_email(req.email)
    if not db_user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    if not verify_password(req.password, db_user["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    if db_user["role"] != req.role:
        raise HTTPException(status_code=403, detail="Incorrect role")

    token = create_access_token(
        user_id=db_user["id"],
        email=db_user["email"],
        role=db_user["role"],
        name=db_user["name"],
    )

    safe_user = {k: v for k, v in db_user.items() if k != "password"}
    response = JSONResponse(content={"user": safe_user})
    response.set_cookie(
        key=JWT_COOKIE_NAME,
        value=token,
        httponly=True,
        samesite="lax",
        secure=False,
        max_age=8 * 3600,
    )
    return response


@app.post("/auth/logout")
async def logout():
    response = JSONResponse({"status": "logged out"})
    response.delete_cookie(JWT_COOKIE_NAME)
    return response


@app.get("/auth/me")
async def get_me(user: dict = Depends(get_current_user)):
    return {"user": user}


# ════════════════════════════════════════════════
# SUBJECTS
# ════════════════════════════════════════════════

@app.post("/subjects")
async def create_subject_endpoint(
    req: CreateSubjectRequest,
    user: dict = Depends(get_current_user)
):
    subject_id = create_subject(
        name=req.name,
        description=req.description,
        created_by=user.get("sub", "")
    )
    return {"id": subject_id, "name": req.name, "description": req.description}


@app.get("/subjects")
async def list_subjects(user: dict = Depends(get_current_user)):
    subjects = get_all_subjects()
    # Attach in-memory status (whether materials are loaded)
    for s in subjects:
        s["is_loaded"] = s["id"] in subject_stores
    return subjects


@app.get("/subjects/{subject_id}")
async def get_subject(subject_id: str, user: dict = Depends(get_current_user)):
    subject = get_subject_by_id(subject_id)
    if not subject:
        raise HTTPException(404, "Subject not found")
    subject["is_loaded"] = subject_id in subject_stores
    return subject


@app.delete("/subjects/{subject_id}")
async def remove_subject(subject_id: str, user: dict = Depends(get_current_user)):
    require_role(user, ["admin"])
    delete_subject(subject_id)
    # Clear in-memory stores
    subject_stores.pop(subject_id, None)
    subject_chunks.pop(subject_id, None)
    return {"status": "deleted"}


# ════════════════════════════════════════════════
# MATERIALS — upload per subject
# ════════════════════════════════════════════════

@app.post("/subjects/{subject_id}/upload")
async def upload_material(
    subject_id: str,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user)
):
    subject = get_subject_by_id(subject_id)
    if not subject:
        raise HTTPException(404, "Subject not found")

    file_id = str(uuid.uuid4())
    file_path = f"{UPLOAD_DIR}/{file_id}_{file.filename}"

    with open(file_path, "wb") as f:
        f.write(await file.read())

    # Parse and index — supports all file formats
    content = load_file(file_path)
    if not content.strip():
        raise HTTPException(400, f"Could not extract text from '{file.filename}'. Check the file is not empty or password-protected.")

    chunks = chunk_text(content)
    if not chunks:
        raise HTTPException(400, "No content chunks found in PDF")

    # Persist to DB
    material_id = insert_material(
        subject_id=subject_id,
        filename=file.filename,
        file_path=file_path,
        uploaded_by=user.get("email", "")
    )

    # Index into in-memory store
    _index_chunks_for_subject(subject_id, chunks)

    return {
        "id": material_id,
        "filename": file.filename,
        "subject_id": subject_id,
        "chunks_indexed": len(chunks),
        "status": "indexed"
    }


@app.get("/subjects/{subject_id}/materials")
async def list_materials(subject_id: str, user: dict = Depends(get_current_user)):
    return get_materials_by_subject(subject_id)


@app.delete("/materials/{material_id}")
async def remove_material(material_id: str, user: dict = Depends(get_current_user)):
    require_role(user, ["admin"])
    delete_material(material_id)
    return {"status": "deleted"}


# ════════════════════════════════════════════════
# PAPERS — history per subject
# ════════════════════════════════════════════════

@app.get("/subjects/{subject_id}/papers")
async def list_papers(subject_id: str, user: dict = Depends(get_current_user)):
    return get_papers_by_subject(subject_id)


@app.get("/papers/{paper_id}")
async def get_paper(paper_id: str, user: dict = Depends(get_current_user)):
    paper = get_paper_with_questions(paper_id)
    if not paper:
        raise HTTPException(404, "Paper not found")
    return paper


# ════════════════════════════════════════════════
# BLOOM / DIFFICULTY HELPERS  (defined before use)
# ════════════════════════════════════════════════

def _bloom_to_verb(bloom: str, question_type: str = "") -> str:
    """Map Bloom's taxonomy level to an appropriate question verb."""
    BLOOM_VERBS = {
        "Remember":   ["Define", "List", "Recall", "Name", "State"],
        "Understand": ["Explain", "Describe", "Summarise", "Interpret", "Classify"],
        "Apply":      ["Apply", "Solve", "Demonstrate", "Calculate", "Use"],
        "Analyse":    ["Analyse", "Compare", "Differentiate", "Examine", "Investigate"],
        "Evaluate":   ["Evaluate", "Justify", "Critique", "Assess", "Argue"],
        "Create":     ["Design", "Construct", "Propose", "Develop", "Formulate"],
    }
    return random.choice(BLOOM_VERBS.get(bloom, ["Explain"]))


def _difficulty_to_focus(difficulty: str) -> str:
    FOCUS = {
        "Easy":   "basic conceptual understanding",
        "Medium": "applied and analytical thinking",
        "Hard":   "critical evaluation and deep insight",
    }
    return FOCUS.get(difficulty, "conceptual understanding")


# ════════════════════════════════════════════════
# GENERATE PAPER (SSE STREAM)
#
# Pattern format expected from frontend:
# {
#   "subject_id": "...",
#   "title": "Mid-Sem Exam",
#   "pattern": {
#     "sections": [
#       {
#         "name": "A",
#         "type": "MCQ",          ← question type
#         "count": 10,
#         "marksPerQuestion": 1,
#         "bloomLevel": "Remember",
#         "difficulty": "Easy"
#       },
#       {
#         "name": "B",
#         "type": "Short Answer",
#         "count": 5,
#         "marksPerQuestion": 3,
#         "bloomLevel": "Understand",
#         "difficulty": "Medium"
#       }
#     ]
#   }
# }
# ════════════════════════════════════════════════

@app.post("/generate/paper")
async def generate_paper(
    req: GeneratePaperRequest,
    user: dict = Depends(get_current_user)
):
    async def event_stream():
        try:
            subject_id = req.subject_id

            # Validate subject
            subject = get_subject_by_id(subject_id)
            if not subject:
                yield f"data: {json.dumps({'type':'error','message':'Subject not found'})}\n\n"
                yield f"data: {json.dumps({'type':'done'})}\n\n"
                return

            # Validate materials loaded
            chunks = subject_chunks.get(subject_id, [])
            store = subject_stores.get(subject_id)

            if not chunks or not store:
                yield f"data: {json.dumps({'type':'error','message':'No materials uploaded for this subject. Please upload PDFs first.'})}\n\n"
                yield f"data: {json.dumps({'type':'done'})}\n\n"
                return

            # Create paper record
            paper_id = create_paper(
                subject_id=subject_id,
                created_by=user.get("email", ""),
                title=req.title or "Untitled Paper"
            )

            paper_msg = "Paper created for subject: " + subject["name"]
            yield f"data: {json.dumps({'type':'step','message':paper_msg})}\n\n"

            sections_response = []

            for section in req.pattern.get("sections", []):
                section_name   = section.get("name", "A")
                count          = int(section.get("count", 1))
                marks          = int(section.get("marksPerQuestion", 2))
                question_type  = section.get("type", "Short Answer")
                # ── NOW READS FROM PATTERN (not hardcoded) ──
                bloom          = section.get("bloomLevel", "Understand")
                difficulty     = section.get("difficulty", "Medium")

                section_questions = []

                # Generate queries once per section (not per question — avoids repeated model calls)
                section_queries = query_generator.generate_queries(chunks, top_k=max(count, 5))
                if not section_queries:
                    section_queries = [bloom]

                for i in range(count):
                    step_msg = f"Section {section_name} Q{i+1}: retrieving context..."
                    yield f"data: {json.dumps({'type':'step','message':step_msg})}\n\n"

                    # Rotate through pre-generated queries for variety
                    query = section_queries[i % len(section_queries)]

                    context_chunks = store.search(query, top_k=5)

                    # Map bloom level to a suitable verb
                    verb = _bloom_to_verb(bloom, question_type)

                    prompt = build_prompt(
                        context_chunks=context_chunks,
                        question_type=question_type,
                        bloom_level=bloom,
                        difficulty=difficulty,
                        marks=marks,
                        verb=verb,
                        focus=_difficulty_to_focus(difficulty)
                    )

                    question_text = generate_question(prompt)

                    # Similarity check — one retry if similar
                    if is_similar(question_text):
                        regen_msg = f"Section {section_name} Q{i+1}: similar question detected, regenerating..."
                        yield f"data: {json.dumps({'type':'step','message':regen_msg})}\n\n"
                        question_text = generate_question(prompt)

                    question_id = insert_question(
                        paper_id=paper_id,
                        section_name=section_name,
                        question_text=question_text,
                        bloom_level=bloom,
                        difficulty=difficulty,
                        marks=marks,
                        similarity=0,
                        source=query,
                        question_type=question_type
                    )

                    section_questions.append({
                        "id": question_id,
                        "text": question_text,
                        "marks": marks,
                        "bloom": bloom,
                        "difficulty": difficulty,
                        "similarity": 0,
                        "source": query,
                        "question_type": question_type
                    })

                sections_response.append({
                    "name": section_name,
                    "questions": section_questions
                })

            yield f"data: {json.dumps({'type':'done','payload':{'paper_id': paper_id,'sections': sections_response}})}\n\n"

        except Exception as e:
            import traceback
            print("STREAM ERROR:", traceback.format_exc())
            yield f"data: {json.dumps({'type':'error','message':str(e)})}\n\n"
            yield f"data: {json.dumps({'type':'done'})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ════════════════════════════════════════════════
# QUESTION REVIEW — edit / regenerate / delete
# ════════════════════════════════════════════════

@app.post("/question/edit")
async def edit_question(req: EditQuestionRequest):
    update_question_text(req.question_id, req.text)
    return {"status": "updated"}


@app.post("/regenerate/question")
async def regenerate_question(
    req: RegenerateRequest,
    user: dict = Depends(get_current_user)
):
    subject_id = req.subject_id
    chunks = subject_chunks.get(subject_id, [])
    store = subject_stores.get(subject_id)

    if not chunks or not store:
        raise HTTPException(400, "No materials loaded for this subject. Upload PDFs first.")

    queries = query_generator.generate_queries(chunks, top_k=3)
    query = random.choice(queries) if queries else "concept"

    context = store.search(query, top_k=5)

    prompt = build_prompt(
        context_chunks=context,
        question_type=req.question_type,
        bloom_level=req.bloom_level,
        difficulty=req.difficulty,
        marks=req.marks,
        verb=_bloom_to_verb(req.bloom_level, req.question_type),
        focus=_difficulty_to_focus(req.difficulty)
    )

    question_text = generate_question(prompt)
    update_question_text(req.question_id, question_text)

    return {"text": question_text}


@app.delete("/question/{question_id}")
async def delete_question(question_id: str):
    delete_question_by_id(question_id)
    return {"status": "deleted"}


# ════════════════════════════════════════════════
# ANALYTICS
# ════════════════════════════════════════════════

@app.get("/analytics/summary")
async def analytics_summary():
    return get_analytics_summary()

@app.get("/analytics/bloom")
async def bloom_dist():
    return get_bloom_distribution()

@app.get("/analytics/difficulty")
async def difficulty_dist():
    return get_difficulty_distribution()


# ════════════════════════════════════════════════
# EXPORT — PDF and DOCX
# ════════════════════════════════════════════════

class ExportRequest(BaseModel):
    paper_id: str


@app.post("/export/pdf")
async def export_pdf(req: ExportRequest, user: dict = Depends(get_current_user)):
    from fpdf import FPDF
    from db import get_paper_with_questions

    paper = get_paper_with_questions(req.paper_id)
    if not paper:
        raise HTTPException(404, "Paper not found")

    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, paper.get("title") or "Question Paper", ln=True, align="C")
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 6, f"Generated: {paper.get('created_at', '')}", ln=True, align="C")
    pdf.ln(6)

    for section in paper.get("sections", []):
        # Section header
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(30, 41, 59)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 8, f"  Section {section['name']}", ln=True, fill=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(3)

        for qi, q in enumerate(section.get("questions", []), 1):
            # Question number + marks
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, f"Q{qi}.  [{q.get('marks', '')} marks | Bloom: {q.get('bloom', '')} | {q.get('difficulty', '')}]", ln=True)

            # Question text — handle unicode by encoding safely
            pdf.set_font("Helvetica", "", 10)
            text = q.get("text", "").encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 6, text)
            pdf.ln(3)

        pdf.ln(4)

    # Stream as response
    import io
    buf = io.BytesIO()
    pdf_bytes = pdf.output()
    buf.write(pdf_bytes)
    buf.seek(0)

    filename = (paper.get("title") or "question_paper").replace(" ", "_") + ".pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@app.post("/export/docx")
async def export_docx(req: ExportRequest, user: dict = Depends(get_current_user)):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from db import get_paper_with_questions
    import io

    paper = get_paper_with_questions(req.paper_id)
    if not paper:
        raise HTTPException(404, "Paper not found")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(paper.get("title") or "Question Paper")
    run.bold = True
    run.font.size = Pt(18)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {paper.get('created_at', '')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    for section in paper.get("sections", []):
        # Section heading
        sec_heading = doc.add_paragraph()
        sec_run = sec_heading.add_run(f"Section {section['name']}")
        sec_run.bold = True
        sec_run.font.size = Pt(13)
        sec_run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

        for qi, q in enumerate(section.get("questions", []), 1):
            # Meta line
            meta = doc.add_paragraph()
            meta_run = meta.add_run(
                f"Q{qi}.  [{q.get('marks', '')} marks | Bloom: {q.get('bloom', '')} | {q.get('difficulty', '')} | {q.get('question_type', '')}]"
            )
            meta_run.bold = True
            meta_run.font.size = Pt(9)
            meta_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

            # Question text
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(q.get("text", ""))
            q_run.font.size = Pt(11)
            q_para.paragraph_format.space_after = Pt(8)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = (paper.get("title") or "question_paper").replace(" ", "_") + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# ════════════════════════════════════════════════
# HEALTH
# ════════════════════════════════════════════════

@app.get("/")
async def health():
    return {
        "status": "QP Backend Running",
        "subjects_loaded": list(subject_stores.keys()),
        "total_chunks": {sid: len(c) for sid, c in subject_chunks.items()}
    }
